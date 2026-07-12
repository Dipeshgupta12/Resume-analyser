from io import BytesIO

from docx import Document


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text content from a DOCX file."""
    document = Document(BytesIO(file_content))
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError("Could not extract text from DOCX. The file may be empty.")
    return text
